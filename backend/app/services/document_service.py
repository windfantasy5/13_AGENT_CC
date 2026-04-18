"""
文档处理服务
"""
import hashlib
import aiofiles
from pathlib import Path
from typing import Optional, Dict, List
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.models.document import Document, DocumentChunk
from app.core.text_splitter import SmartTextSplitter
from app.core.pdf_processor import PDFProcessor
from app.core.vector_store import VectorStore
from app.core.web_scraper import WebScraper
from app.config.settings import settings
import logging
from docx import Document as DocxDocument
import io

logger = logging.getLogger(__name__)


class DocumentService:
    """文档处理服务"""

    def __init__(self):
        self.text_splitter = SmartTextSplitter()
        self.pdf_processor = PDFProcessor()
        self.vector_store = VectorStore()
        self.web_scraper = WebScraper()
        self.upload_dir = Path(settings.upload_dir_abs_path)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def process_file_upload(
        self,
        file_content: bytes,
        filename: str,
        title: str,
        user_id: int,
        db: AsyncSession
    ) -> Document:
        """
        处理文件上传

        Args:
            file_content: 文件内容
            filename: 文件名
            title: 文档标题
            user_id: 用户ID
            db: 数据库会话

        Returns:
            Document对象
        """
        # 计算文件哈希
        file_hash = hashlib.sha256(file_content).hexdigest()

        # 检查是否已存在
        result = await db.execute(
            select(Document).where(Document.file_hash == file_hash)
        )
        existing_doc = result.scalar_one_or_none()
        if existing_doc:
            raise Exception("文档已存在")

        # 判断文件类型
        file_type = self._get_file_type(filename)

        # 保存文件
        file_path = self.upload_dir / f"{file_hash}_{filename}"
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)

        # 创建文档记录
        document = Document(
            user_id=user_id,
            title=title,
            file_type=file_type,
            file_path=str(file_path),
            file_hash=file_hash,
            file_size=len(file_content),
            status='pending'
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)

        return document

    async def extract_text_from_file(
        self,
        document: Document
    ) -> str:
        """
        从文件提取文本

        Args:
            document: 文档对象

        Returns:
            提取的文本
        """
        try:
            # 读取文件
            async with aiofiles.open(document.file_path, 'rb') as f:
                file_content = await f.read()

            # 根据文件类型提取文本
            if document.file_type == 'pdf':
                text, used_ocr = self.pdf_processor.extract_text(file_content)
                if used_ocr:
                    logger.info(f"文档 {document.id} 使用OCR识别")
            elif document.file_type == 'word':
                text = self._extract_from_word(file_content)
            elif document.file_type == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
            else:
                raise Exception(f"不支持的文件类型: {document.file_type}")

            return text

        except Exception as e:
            logger.error(f"提取文本失败: {e}")
            raise

    async def extract_text_from_bytes(
        self,
        file_content: bytes,
        filename: str
    ) -> str:
        """
        从字节内容提取文本

        Args:
            file_content: 文件内容
            filename: 文件名

        Returns:
            提取的文本
        """
        try:
            # 判断文件类型
            file_type = self._get_file_type(filename)

            # 根据文件类型提取文本
            if file_type == 'pdf':
                text, used_ocr = self.pdf_processor.extract_text(file_content)
                if used_ocr:
                    logger.info(f"文件 {filename} 使用OCR识别")
            elif file_type == 'word':
                text = self._extract_from_word(file_content)
            elif file_type == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
            else:
                raise Exception(f"不支持的文件类型: {file_type}")

            return text

        except Exception as e:
            logger.error(f"提取文本失败: {e}")
            raise

    async def extract_text_from_url(self, url: str) -> Dict[str, str]:
        """
        从URL提取文本

        Args:
            url: 网页URL

        Returns:
            {title, content, markdown}
        """
        return await self.web_scraper.fetch_and_extract(url)

    def preview_chunks(
        self,
        text: str,
        params: Optional[Dict] = None,
        page: int = 1,
        page_size: int = 10
    ) -> Dict:
        """
        预览文档分段

        Args:
            text: 文本内容
            params: 分段参数
            page: 页码
            page_size: 每页数量

        Returns:
            分段预览结果
        """
        # 执行分段
        preview_result = self.text_splitter.preview_split(text, params)

        # 返回所有分段（不分页，前端需要看到全部）
        return {
            "chunks": preview_result["chunks"],
            "total_chunks": preview_result["total_chunks"],
            "total_chars": preview_result["total_chars"],
            "avg_chunk_size": preview_result["avg_chunk_size"],
            "params": preview_result["params"]
        }

    async def save_chunks(
        self,
        document_id: int,
        text: str,
        params: Optional[Dict],
        db: AsyncSession
    ) -> int:
        """
        保存文档分段并向量化

        Args:
            document_id: 文档ID
            text: 文本内容
            params: 分段参数
            db: 数据库会话

        Returns:
            分段数量
        """
        # 执行分段
        if params:
            splitter = SmartTextSplitter(
                max_chunk_size=params.get("max_chunk_size", 500),
                min_chunk_size=params.get("min_chunk_size", 50),
                overlap_size=params.get("overlap_size", 50)
            )
            chunks = splitter.split_text(text)
        else:
            chunks = self.text_splitter.split_text(text)

        # 保存分段到数据库
        chunk_records = []
        for chunk_data in chunks:
            chunk = DocumentChunk(
                document_id=document_id,
                chunk_index=chunk_data["index"],
                content=chunk_data["content"],
                char_count=chunk_data["char_count"]
            )
            db.add(chunk)
            chunk_records.append(chunk)

        # 提交以获取chunk ID
        await db.commit()
        for chunk in chunk_records:
            await db.refresh(chunk)

        # 向量化并存储
        try:
            chunk_ids = [chunk.id for chunk in chunk_records]
            contents = [chunk.content for chunk in chunk_records]
            metadatas = [
                {
                    "document_id": document_id,
                    "chunk_index": chunk.chunk_index,
                    "char_count": chunk.char_count
                }
                for chunk in chunk_records
            ]

            # 添加到向量库
            vector_ids = self.vector_store.add_chunks(
                chunk_ids=chunk_ids,
                contents=contents,
                metadatas=metadatas
            )

            # 更新vector_id
            for i, chunk in enumerate(chunk_records):
                chunk.vector_id = vector_ids[i]

            logger.info(f"文档 {document_id} 的 {len(chunks)} 个分块已向量化")

        except Exception as e:
            logger.error(f"向量化失败: {e}")
            # 向量化失败不影响分块保存，继续执行

        # 更新文档状态
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one()
        document.chunk_count = len(chunks)
        document.status = 'completed'

        await db.commit()

        return len(chunks)

    def _get_file_type(self, filename: str) -> str:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext in ['.doc', '.docx']:
            return 'word'
        elif ext == '.txt':
            return 'txt'
        else:
            raise Exception(f"不支持的文件格式: {ext}")

    def _extract_from_word(self, file_content: bytes) -> str:
        """从Word文档提取文本"""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return '\n\n'.join(paragraphs)
        except Exception as e:
            logger.error(f"提取Word文本失败: {e}")
            raise Exception("Word文档解析失败")
